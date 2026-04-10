"""
qa_agent.py — 사업계획서 QA 검증 에이전트
==========================================
역할: 완성된 DOCX와 content를 비교하여 오류·경고 목록 반환.

검사 항목:
    E01~E0x: 필수 셀 빈 값
    W01~W0x: 경고 (이미지 수 초과, 비중 불일치 등)
    W05: 이미지 4장 초과 (데이터바우처 규정: 1장)

TODO 추가 검사 항목:
    - DOCX 내 ○○○ 텍스트 잔존 (공란 미입력)
    - schedule weight 합계 ≠ 100
    - budget_mix 합계 ≠ 100
    - 팀 참여율 합계 > 100%

반환:
    {"error_count": int, "warning_count": int, "errors": [...], "warnings": [...]}
"""
from __future__ import annotations
import os
import re

from docx import Document


# 과제명에 반드시 포함되어야 할 데이터 관련 단어들
DATA_KEYWORDS = [
    "데이터", "data", "DB", "분석", "AI", "머신러닝", "딥러닝",
    "빅데이터", "크롤링", "수집", "정형화", "가공"
]


class QAAgent:
    """
    QA 에이전트 — 최종 제출 가능 여부 판정.

    run(content, output_docx_path) → QAReport dict
    """

    def __init__(self):
        self.errors:   list[str] = []
        self.warnings: list[str] = []

    # ── 메인 실행 ──────────────────────────────────────────────
    def run(self, content: dict, output_path: str = "") -> dict:
        """
        content 전체를 순회하며 데이터바우처 규정 위반 사항 검사.

        Args:
            content:     통합 content (모든 에이전트 처리 완료)
            output_path: 생성된 DOCX 경로 (파일 존재 여부 확인용)

        Returns:
            QA 리포트 딕셔너리
        """
        self.errors   = []
        self.warnings = []

        self._check_errors(content)
        self._check_warnings(content)
        self._collect_agent_issues(content)
        self._check_output_docx(output_path)

        # 출력 파일 존재 여부
        # dry_run=True(BizPlanInjector 미설치) 상태에서는 .docx가 생성되지 않으므로 검사 skip
        is_dry_run = content.get("_render_stats", {}).get("dry_run", False)
        if output_path and not os.path.exists(output_path) and not is_dry_run:
            self.errors.append(f"출력 파일이 생성되지 않았습니다: {output_path}")

        return {
            "passed":        len(self.errors) == 0,
            "error_count":   len(self.errors),
            "warning_count": len(self.warnings),
            "errors":        self.errors,
            "warnings":      self.warnings,
        }

    def _check_output_docx(self, output_path: str):
        if not output_path or not os.path.exists(output_path):
            return

        try:
            doc = Document(output_path)
        except Exception as exc:
            self.errors.append(f"[E07] 출력 DOCX를 읽을 수 없습니다: {exc}")
            return

        blocker_patterns = {
            "[E08] DOCX에 미확정 placeholder '○○○'가 남아 있습니다.": "○○○",
            "[E09] DOCX에 미입력 직급 placeholder '…'가 남아 있습니다.": "…",
            "[E10] DOCX에 이미지 가이드 문구가 남아 있습니다.": "[이미지 삽입 위치]",
            "[E11] DOCX에 template 입력 안내 문구가 남아 있습니다.": "project_input.json",
            "[E12] DOCX에 작성요령 또는 가이드 표 문구가 남아 있습니다.": "< 작성요령 >",
        }

        all_text = []
        for para in doc.paragraphs:
            if para.text:
                all_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        all_text.append(cell.text)
        merged_text = "\n".join(all_text)

        for message, token in blocker_patterns.items():
            if token in merged_text and message not in self.errors:
                self.errors.append(message)

        if "라. 개인정보 관리방안" in merged_text and "해당없음" not in merged_text and not re.search(r"개인정보보호관리자", merged_text):
            self.warnings.append(
                "[W08] 개인정보 관리방안 섹션에 실제 내용 또는 '해당없음' 표시가 없습니다."
            )

    # ── 필수 오류 검사 ─────────────────────────────────────────
    def _check_errors(self, content: dict):
        meta      = content.get("meta", {})
        budget    = content.get("budget_mix", {})
        narrative = content.get("narrative", {})
        team      = content.get("team", [])
        privacy   = content.get("privacy", {})

        # E01: 과제명 데이터 키워드
        title = meta.get("project_title", "")
        title_lower = title.lower()
        if not any(kw.lower() in title_lower for kw in DATA_KEYWORDS):
            self.errors.append(
                f"[E01] 과제명에 데이터 관련 단어가 없습니다: '{title}'\n"
                f"      데이터바우처는 데이터 중심 과제명이 필수입니다."
            )

        # E02: 구매/가공 필수
        if budget.get("purchase", 0) == 0 and budget.get("processing", 0) == 0:
            self.errors.append(
                "[E02] 사업비 편성에 구매(purchase) 또는 가공(processing) 항목이 모두 0%입니다.\n"
                "      데이터바우처는 구매 또는 가공 중 최소 1개가 반드시 포함되어야 합니다."
            )

        # E03: 사업비 합계
        total = sum(budget.get(k, 0) for k in
                    ["planning", "purchase", "collection", "processing", "analysis"])
        if total != 100:
            self.errors.append(
                f"[E03] 사업비 비중 합계가 {total}%입니다 (100% 필요).\n"
                f"      현재: planning={budget.get('planning',0)} / "
                f"purchase={budget.get('purchase',0)} / "
                f"collection={budget.get('collection',0)} / "
                f"processing={budget.get('processing',0)} / "
                f"analysis={budget.get('analysis',0)}"
            )

        # E04: 개인정보 관리자
        if privacy.get("uses_personal_data", False):
            if not privacy.get("privacy_manager", "").strip():
                self.errors.append(
                    "[E04] 개인정보 활용 과제이나 개인정보보호관리자가 지정되지 않았습니다.\n"
                    "      privacy.privacy_manager 필드를 채워주세요."
                )

        # E05: 수행인력 최소 1명
        if not team:
            self.errors.append(
                "[E05] 수행인력이 입력되지 않았습니다. 최소 1명 이상 필요합니다."
            )

        # E06: 핵심 서술 섹션 누락
        required_narrative = ["problem", "solution", "market_status"]
        for field in required_narrative:
            val = narrative.get(field, "").strip()
            if not val or "[" in val[:20]:  # 템플릿 placeholder만 있는 경우
                self.errors.append(
                    f"[E06] 서술 섹션 '{field}'이(가) 비어 있거나 초안 상태입니다.\n"
                    f"      실제 내용을 project_input.json의 narrative.{field}에 작성해주세요."
                )

    # ── 경고 검사 ──────────────────────────────────────────────
    def _check_warnings(self, content: dict):
        meta      = content.get("meta", {})
        budget    = content.get("budget_mix", {})
        narrative = content.get("narrative", {})
        team      = content.get("team", [])
        schedule  = content.get("schedule", [])

        # W01: goal_100char 길이
        goal = meta.get("goal_100char", "")
        if len(goal) > 100:
            self.warnings.append(
                f"[W01] 데이터 활용 목표가 {len(goal)}자입니다 (100자 이내 권장)."
            )

        # W02: 단일 항목 편중
        max_val = max(budget.values(), default=0)
        if max_val > 70:
            self.warnings.append(
                f"[W02] 사업비 단일 항목이 {max_val}%로 과도합니다. "
                "70% 초과 시 미선정 위험이 있습니다."
            )

        # W03: 일정 비중 합계
        sched_total = sum(s.get("weight", 0) for s in schedule)
        if schedule and sched_total != 100:
            self.warnings.append(
                f"[W03] 추진일정 비중 합계가 {sched_total}%입니다 (100% 권장)."
            )

        # W04: 참여율 범위
        for m in team:
            p = m.get("participation", 0)
            if not (10 <= p <= 100):
                self.warnings.append(
                    f"[W04] '{m.get('name','')}' 참여율 {p}%가 권고 범위(10~100%)를 벗어납니다."
                )

        # W05: 이미지 경고 (AssetAgent에서 수집된 것)
        for w in content.get("_asset_warnings", []):
            self.warnings.append(f"[W05] {w}")

        # W06: 과제명 길이
        title = meta.get("project_title", "")
        if len(title) > 50:
            self.warnings.append(
                f"[W06] 과제명이 {len(title)}자입니다. 50자 이하를 권장합니다."
            )

        # W07: expected_effect 누락
        ee = narrative.get("expected_effect", "").strip()
        if not ee:
            self.warnings.append(
                "[W07] 기대효과(expected_effect) 항목이 비어 있습니다. 작성을 권장합니다."
            )

    # ── 에이전트 내부 이슈 수집 ───────────────────────────────
    def _collect_agent_issues(self, content: dict):
        """Table / Asset 에이전트가 수집한 오류/경고를 QA 리포트에 통합."""
        for e in content.get("_table_errors", []):
            if e not in self.errors:
                self.errors.append(f"[TableAgent] {e}")
        for w in content.get("_table_warnings", []):
            if w not in self.warnings:
                self.warnings.append(f"[TableAgent] {w}")
        for e in content.get("_asset_errors", []):
            if e not in self.errors:
                self.errors.append(f"[AssetAgent] {e}")

    # ── 상태 요약 ──────────────────────────────────────────────
    def summary(self) -> str:
        return (f"errors={len(self.errors)}, warnings={len(self.warnings)}, "
                f"passed={len(self.errors)==0}")
