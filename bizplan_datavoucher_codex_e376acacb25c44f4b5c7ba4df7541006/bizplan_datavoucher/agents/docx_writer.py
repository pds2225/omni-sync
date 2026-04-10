"""
docx_writer.py — python-docx 기반 DOCX 직접 조작 엔진
=======================================================
역할: render_agent._direct_write()에서 호출되는 실제 DOCX 생성 로직.
      템플릿 DOCX를 복사 후 injector_content의 데이터를 삽입·삭제·서식화.

진입점:
    write_docx(template_path, output_path, injector_content) -> stats_dict

실행 순서 (write_docx 내부):
    STEP 1   셀 채우기           (injector_content["cells"])
    STEP 1-B 잉여 빈 행 제거     (injector_content["trim_tables"])
    STEP 1-C 실제 이미지 삽입    (injector_content["image_slots"])
    STEP 2   섹션 텍스트 삽입    (injector_content["sections"])
    STEP 2-B 이미지 placeholder  (injector_content["image_placeholders"])
    STEP 2-C 빈 칸 표 삽입       (injector_content["blank_image_slots"])
    STEP 3   작성요령 삭제        (delete_guide_elements)
    STEP 4   연속 공백 압축       (collapse_blank_paras, max_blank=1)

색상 초기화 정책:
    _should_reset_color(val) → _PRESERVE_COLORS(흰색 계열)는 보존, 나머지는 000000

⚠️ 의존성: python-docx, lxml
⚠️ TABLE_INDEX는 작성요령 표 삭제 후 기준 (render_agent.py 참조)
"""
import copy
import os
import shutil
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt

# ── 색상 초기화 정책: 흰색(ffffff) 등 배경 대비 텍스트 색상은 보존 ──────────
_PRESERVE_COLORS = {
    "ffffff", "fffffe", "fefefe",   # 흰색 (어두운 배경 위 텍스트)
    "f2f2f2", "f0f0f0",             # 밝은 회색
}

def _should_reset_color(val: str) -> bool:
    """색상을 000000으로 초기화해야 하면 True."""
    if not val:
        return False
    return val.lower() not in _PRESERVE_COLORS

from lxml import etree

# ── 색상 감지 ────────────────────────────────────────────────────
_BLUE_COLORS = {
    "4472c4", "1f3864", "2e74b5", "4f81bd", "17375e", "244185",
    "1f497d", "0070c0", "2f5496", "215868", "1f5c8b", "0000ff",
    "002060", "305496",
}

def _is_blue_rpr(rpr) -> bool:
    if rpr is None: return False
    col = rpr.find(qn("w:color"))
    if col is None: return False
    return col.get(qn("w:val"), "").lower() in _BLUE_COLORS

def _is_blue_para(p_elem) -> bool:
    for r in p_elem.findall(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if _is_blue_rpr(rpr): return True
    ppr = p_elem.find(qn("w:pPr"))
    if ppr is not None:
        rpr = ppr.find(qn("w:rPr"))
        if _is_blue_rpr(rpr): return True
    return False

def _tbl_has_blue(tbl_elem) -> bool:
    for p in tbl_elem.iter(qn("w:p")):
        if _is_blue_para(p): return True
    return False

def _tbl_is_guide(tbl_elem) -> bool:
    """
    표가 < 작성요령 > 안내 표인지 판별.
    - 어느 셀이든 '< 작성요령 >' 텍스트가 포함되면 True.
    - 파란색 여부와 무관하게 텍스트 기반으로 판별 (서식이 초기화된 경우도 대응).
    """
    for tc in tbl_elem.iter(qn("w:tc")):
        t = "".join(x.text or "" for x in tc.iter(qn("w:t"))).strip()
        if "< 작성요령 >" in t or "<작성요령>" in t:
            return True
    return False


# ── 셀 조작 ──────────────────────────────────────────────────────
def _clear_cell(tc):
    """셀의 모든 단락 텍스트를 비우고, 파란색·이탤릭 서식도 초기화."""
    for p in tc.findall(qn("w:p")):
        # pPr > rPr 색상도 초기화 (run 기본 서식 상속 차단)
        ppr = p.find(qn("w:pPr"))
        if ppr is not None:
            para_rpr = ppr.find(qn("w:rPr"))
            if para_rpr is not None:
                col = para_rpr.find(qn("w:color"))
                if col is not None and _should_reset_color(col.get(qn("w:val"), "")):
                    col.set(qn("w:val"), "000000")
                for itag in [qn("w:i"), qn("w:iCs")]:
                    el = para_rpr.find(itag)
                    if el is not None: para_rpr.remove(el)
        for r in p.findall(qn("w:r")):
            for t in r.findall(qn("w:t")):
                t.text = ""
            # 서식 초기화: 파란색 → 검정, 이탤릭 제거
            rpr = r.find(qn("w:rPr"))
            if rpr is not None:
                col = rpr.find(qn("w:color"))
                cur_val = col.get(qn("w:val"), "") if col is not None else ""
                if _should_reset_color(cur_val):
                    if col is None:
                        col = etree.SubElement(rpr, qn("w:color"))
                    col.set(qn("w:val"), "000000")
                for tag in [qn("w:i"), qn("w:iCs")]:
                    el = rpr.find(tag)
                    if el is not None: rpr.remove(el)

def _set_cell_text(tc, text: str):
    """셀에 텍스트를 채움. 기존 내용(안내문 포함) 전체 교체.

    셀 안에 중첩 <w:tbl>(작성요령 안내 표)이 있으면 함께 삭제.
    """
    # ── 중첩 표(작성요령 안내 표) 삭제 ────────────────────────
    for nested_tbl in tc.findall(qn("w:tbl")):
        tc.remove(nested_tbl)

    # ── pPr > rPr 의 파란색·이탤릭 서식 초기화 ─────────────────
    # pPr rPr 의 color 는 단락 내 run 의 기본 서식으로 상속되므로 반드시 초기화
    for p in tc.findall(qn("w:p")):
        ppr = p.find(qn("w:pPr"))
        if ppr is not None:
            para_rpr = ppr.find(qn("w:rPr"))
            if para_rpr is not None:
                col = para_rpr.find(qn("w:color"))
                if col is not None and _should_reset_color(col.get(qn("w:val"), "")):
                    col.set(qn("w:val"), "000000")
                for itag in [qn("w:i"), qn("w:iCs")]:
                    el = para_rpr.find(itag)
                    if el is not None: para_rpr.remove(el)

    paras = tc.findall(qn("w:p"))
    if not paras: return

    lines = text.split("\n")

    # ── 첫 번째 단락: 데이터 삽입 ──────────────────────────────
    _clear_cell(tc)   # 텍스트+서식 초기화
    _write_para(paras[0], lines[0])

    # ── 두 번째~ 단락: 여분 단락 제거 후 새 줄 추가 ────────────
    # 기존 단락이 2개 이상이면 첫 번째 제외 모두 제거 (안내문 단락 삭제)
    for extra_p in paras[1:]:
        tc.remove(extra_p)

    # 새로운 줄 추가 (lines[1:])
    ref_para = paras[0]
    for line in lines[1:]:
        new_p = copy.deepcopy(ref_para)
        for r in new_p.findall(qn("w:r")):
            for t in r.findall(qn("w:t")): t.text = ""
        _write_para(new_p, line)
        tc.append(new_p)

def _write_para(para, text: str):
    """단락에 텍스트를 쓰고, 색상·이탤릭 서식을 검정/일반으로 초기화."""
    runs = para.findall(qn("w:r"))
    if runs:
        for r in runs:
            for t in r.findall(qn("w:t")): t.text = ""
            # ── 서식 초기화: 색상을 검정, 이탤릭 제거 ──────────────
            rpr = r.find(qn("w:rPr"))
            if rpr is None:
                rpr = etree.SubElement(r, qn("w:rPr"))
            # 색상 → 검정(000000) 으로 고정
            col = rpr.find(qn("w:color"))
            if col is None:
                col = etree.SubElement(rpr, qn("w:color"))
            col.set(qn("w:val"), "000000")
            # 이탤릭 제거
            for tag in [qn("w:i"), qn("w:iCs")]:
                el = rpr.find(tag)
                if el is not None: rpr.remove(el)
        t_elems = runs[0].findall(qn("w:t"))
        if t_elems:
            t_elems[0].text = text
            if " " in text:
                t_elems[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        else:
            t_e = etree.SubElement(runs[0], qn("w:t"))
            t_e.text = text
    else:
        r_e = etree.SubElement(para, qn("w:r"))
        t_e = etree.SubElement(r_e, qn("w:t"))
        t_e.text = text
        if " " in text:
            t_e.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


# ── 섹션 삽입 ────────────────────────────────────────────────────
def _para_is_blank(p_elem) -> bool:
    """단락이 완전히 비어있는지 확인.
    ⚠️  페이지 브레이크(w:br type='page'|'column'|'evenPage'|'oddPage') 단락은
        텍스트가 없어도 '빈 단락'으로 처리하지 않음 (레이아웃 필수).
    """
    # 페이지 브레이크 단락 보존
    for br in p_elem.iter(qn("w:br")):
        br_type = br.get(qn("w:type"), "")
        if br_type in ("page", "column", "evenPage", "oddPage"):
            return False
    return not "".join(t.text or "" for t in p_elem.iter(qn("w:t"))).strip()

def _find_section_para(doc: Document, keyword: str):
    for para in doc.paragraphs:
        if keyword in para.text:
            return para
    return None

def _insert_text_after_para(para, text: str, doc: Document) -> bool:
    """
    keyword 단락 바로 다음에 텍스트를 삽입.
    
    핵심: ref_para로 비어있지 않은 첫 번째 단락을 사용하여
    원본의 공백 단락이 대량 복제되는 현상을 방지.
    """
    body = doc.element.body
    try:
        insert_idx = list(body).index(para._element)
    except ValueError:
        return False

    # ── ref_para: 내용이 있는 단락 참조용 (공백 단락 참조 차단) ──
    # 현재 keyword 단락 자체를 ref로 사용
    ref_para = para._element

    lines = text.split("\n")
    for offset, line in enumerate(lines):
        new_p = copy.deepcopy(ref_para)

        # 모든 텍스트 run 비우기
        for r in new_p.findall(qn("w:r")):
            for t in r.findall(qn("w:t")): t.text = ""

        # 파란색/이탤릭 서식 제거 (일반 텍스트로)
        for r in new_p.findall(qn("w:r")):
            rpr = r.find(qn("w:rPr"))
            if rpr is not None:
                for tag in [qn("w:color"), qn("w:i"), qn("w:iCs"),
                             qn("w:u"), qn("w:strike")]:
                    el = rpr.find(tag)
                    if el is not None: rpr.remove(el)

        # 들여쓰기 제거 (헤딩 단락 서식 상속 방지)
        ppr = new_p.find(qn("w:pPr"))
        if ppr is not None:
            for tag in [qn("w:ind"), qn("w:jc"), qn("w:outlineLvl"),
                         qn("w:pStyle")]:
                el = ppr.find(tag)
                if el is not None: ppr.remove(el)
            # pPr 내 rPr 색상 제거
            rpr = ppr.find(qn("w:rPr"))
            if rpr is not None:
                for tag in [qn("w:color"), qn("w:i"), qn("w:iCs")]:
                    el = rpr.find(tag)
                    if el is not None: rpr.remove(el)

        _write_para(new_p, line)
        body.insert(insert_idx + 1 + offset, new_p)

    return True


def _shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_vertical_center(cell):
    try:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    except Exception:
        pass


def _set_table_cell_text(cell, text: str, *, bold: bool = False, align: str = "center",
                         font_size_pt: int = 10):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size_pt)
    _set_cell_vertical_center(cell)


def insert_table_after_para(doc: Document,
                            keyword: str,
                            title: str,
                            headers: list[str],
                            rows: list[list[str]],
                            note: str = "",
                            insert_offset: int = 1,
                            column_widths_cm: list[float] | None = None) -> bool:
    """keyword 단락 뒤에 제출용 실제 DOCX 표를 삽입."""
    para = _find_section_para(doc, keyword)
    if para is None:
        return False

    body = doc.element.body
    try:
        insert_idx = list(body).index(para._element)
    except ValueError:
        return False

    cursor = insert_idx + insert_offset
    ref_para = para._element

    if title:
        title_para = copy.deepcopy(ref_para)
        for r in title_para.findall(qn("w:r")):
            for t in r.findall(qn("w:t")):
                t.text = ""
        _write_para(title_para, title)
        for r in title_para.findall(qn("w:r")):
            rpr = r.find(qn("w:rPr"))
            if rpr is None:
                rpr = etree.SubElement(r, qn("w:rPr"))
            b = rpr.find(qn("w:b"))
            if b is None:
                rpr.append(OxmlElement("w:b"))
            bcs = rpr.find(qn("w:bCs"))
            if bcs is None:
                rpr.append(OxmlElement("w:bCs"))
        body.insert(cursor, title_para)
        cursor += 1

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        _set_table_cell_text(cell, header, bold=True, align="center")
        _shade_cell(cell, "D9EAF7")
        if column_widths_cm and col_idx < len(column_widths_cm):
            cell.width = Cm(column_widths_cm[col_idx])

    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_values):
            cell = table.rows[row_idx].cells[col_idx]
            align = "left" if col_idx == 0 else "center"
            _set_table_cell_text(cell, value, align=align)
            if column_widths_cm and col_idx < len(column_widths_cm):
                cell.width = Cm(column_widths_cm[col_idx])

    tbl_elem = table._tbl
    body.remove(tbl_elem)
    body.insert(cursor, tbl_elem)
    cursor += 1

    if note:
        note_para = copy.deepcopy(ref_para)
        for r in note_para.findall(qn("w:r")):
            for t in r.findall(qn("w:t")):
                t.text = ""
        _write_para(note_para, note)
        body.insert(cursor, note_para)

    return True


# ── 작성요령 / 파란색 삭제 ────────────────────────────────────────
# ─── 내용 기반 표 분류 헬퍼 ──────────────────────────────────────
# (더 이상 인덱스 기반 _KEEP_TABLES를 쓰지 않음 — 삽입 순서 변동에 무관)
_KEEP_TABLES = {4, 13, 22, 27, 29, 30}  # 하위 호환 참조용 (미사용)

def _clean_data_table_inplace(tbl_elem) -> int:
    """
    데이터 보존 표 내부에서:
      1. 행 전체가 파란 이탤릭  → 행 전체 삭제 (안내용 행)
      2. 개별 셀이 파란 이탤릭  → 셀 텍스트 초기화
      3. 모든 run의 파란색/이탤릭 서식  → 검정/일반으로 초기화
    반환값: 삭제된 행 수
    """
    rows_to_del = []
    for tr in tbl_elem.findall(qn("w:tr")):
        cells_in_row = tr.findall(qn("w:tc"))
        has_content = False
        all_blue    = True
        blue_cells  = []

        for tc in cells_in_row:
            tc_has_content = False
            tc_all_blue    = True
            for r in tc.iter(qn("w:r")):
                txt = "".join(t.text or "" for t in r.findall(qn("w:t"))).strip()
                if not txt: continue
                tc_has_content = True
                has_content    = True
                rpr = r.find(qn("w:rPr"))
                if rpr is None:
                    tc_all_blue = False
                    all_blue    = False
                    continue
                col    = rpr.find(qn("w:color"))
                italic = rpr.find(qn("w:i"))
                c_val  = (col.get(qn("w:val"), "") if col is not None else "").lower()
                is_bi  = c_val in ("0000ff", "4472c4", "2e75b6") and italic is not None
                if not is_bi:
                    tc_all_blue = False
                    all_blue    = False
            if tc_has_content and tc_all_blue:
                blue_cells.append(tc)

        if has_content and all_blue:
            rows_to_del.append(tr)
        else:
            for tc in blue_cells:
                _clear_cell(tc)
            # 모든 run 서식 초기화 (run rPr + pPr rPr 모두)
            for tc in cells_in_row:
                # ── pPr > rPr 색상 초기화 (빈 셀 포함) ─────────────
                for p in tc.findall(qn("w:p")):
                    ppr = p.find(qn("w:pPr"))
                    if ppr is not None:
                        para_rpr = ppr.find(qn("w:rPr"))
                        if para_rpr is not None:
                            col = para_rpr.find(qn("w:color"))
                            if col is not None and _should_reset_color(col.get(qn("w:val"), "")):
                                col.set(qn("w:val"), "000000")
                            for itag in [qn("w:i"), qn("w:iCs")]:
                                el = para_rpr.find(itag)
                                if el is not None: para_rpr.remove(el)
                # ── run rPr 색상 초기화 ──────────────────────────────
                for r in tc.iter(qn("w:r")):
                    txt = "".join(t.text or "" for t in r.findall(qn("w:t"))).strip()
                    if not txt: continue
                    rpr = r.find(qn("w:rPr"))
                    if rpr is None: continue
                    col = rpr.find(qn("w:color"))
                    if col is not None and _should_reset_color(col.get(qn("w:val"), "")):
                        col.set(qn("w:val"), "000000")
                    for tag_name in [qn("w:i"), qn("w:iCs")]:
                        el = rpr.find(tag_name)
                        if el is not None: rpr.remove(el)

    deleted = 0
    for tr in rows_to_del:
        try:
            tbl_elem.remove(tr)
            deleted += 1
        except ValueError:
            pass
    return deleted


def delete_guide_elements(doc: Document) -> dict:
    """
    DOCX에서 다음을 삭제:
      1. < 작성요령 > 표  ← 내용 기반 탐지 (인덱스·색상 무관)
      2. 파란색 단락 (제출 시 삭제 안내문구)
      3. 데이터 표 안의 파란 이탤릭 행/셀 (안내용 행 삭제 + 서식 초기화)

    ⚠️  반드시 write_docx()의 셀/섹션 삽입 이후에 호출해야 함.
        내용 기반 판별이므로 빈칸 표·이미지 삽입 후 인덱스 변동과 무관.
    """
    body          = doc.element.body
    tbl_deleted   = 0
    para_deleted  = 0
    blue_row_deleted = 0
    to_remove     = []

    for child in list(body):
        if child.tag == qn("w:tbl"):
            if _tbl_is_guide(child):
                # ── 작성요령 표 → 통째로 삭제 ─────────────────
                to_remove.append(child)
                tbl_deleted += 1
            else:
                # ── 데이터 표 → 행/셀/서식 정리 ─────────────
                blue_row_deleted += _clean_data_table_inplace(child)

        elif child.tag == qn("w:p"):
            if _is_blue_para(child):
                to_remove.append(child)
                para_deleted += 1

    for elem in to_remove:
        try:
            body.remove(elem)
        except ValueError:
            pass

    return {
        "tables_deleted":    tbl_deleted,
        "paras_deleted":     para_deleted,
        "blue_rows_deleted": blue_row_deleted,
    }


# ── 연속 공백 단락 압축 ───────────────────────────────────────────
def collapse_blank_paras(doc: Document, max_blank: int = 1) -> int:
    """
    연속으로 이어지는 빈 단락을 max_blank개 이하로 압축.
    body 직접 자식 단락뿐만 아니라 표 안 셀 내부 단락도 처리.

    Args:
        doc:       대상 Document
        max_blank: 허용할 최대 연속 빈 단락 수 (기본 1)

    Returns:
        삭제한 빈 단락 총 수
    """
    total_deleted = 0

    def _compress(parent_elem, is_cell: bool = False):
        """parent_elem의 직접 자식 중 연속 빈 단락을 압축.
        
        is_cell=True일 때:
          - Word 규격상 마지막 <w:p>는 반드시 보존 (삭제 금지)
          - 단, 그 이전 빈줄들은 max_blank 규칙으로 제거
        """
        deleted = 0
        blank_run = 0
        to_remove = []

        children = list(parent_elem)

        # 셀의 마지막 단락 인덱스(Word 필수 요소) 파악
        last_p_idx = -1
        if is_cell:
            for idx, child in enumerate(children):
                if child.tag == qn("w:p"):
                    last_p_idx = idx

        for idx, child in enumerate(children):
            if child.tag == qn("w:p"):
                # 셀의 마지막 단락은 건드리지 않음 (단, 이전 빈줄 누적 계산은 계속)
                if is_cell and idx == last_p_idx:
                    continue
                is_blank = _para_is_blank(child)
                if is_blank:
                    blank_run += 1
                    if blank_run > max_blank:
                        to_remove.append(child)
                        deleted += 1
                else:
                    blank_run = 0
            elif child.tag == qn("w:tbl"):
                # 표(tbl)가 있으면 공백 연속 카운트를 리셋
                # → 표 삭제 후 앞뒤 공백이 합쳐지는 문제 방지
                blank_run = 0
            else:
                blank_run = 0
        for elem in to_remove:
            try: parent_elem.remove(elem)
            except ValueError: pass
        return deleted

    # 1. body 직접 자식
    total_deleted += _compress(doc.element.body)

    # 2. 모든 표 → 모든 행 → 모든 셀 안 (is_cell=True로 마지막 단락 보존)
    for tbl in doc.element.body.iter(qn("w:tbl")):
        for tc in tbl.iter(qn("w:tc")):
            total_deleted += _compress(tc, is_cell=True)

    return total_deleted


# ── 이미지 Placeholder 삽입 ─────────────────────────────────────

def _insert_image_placeholder_in_cell(doc: Document, table_idx: int,
                                       row: int, cell: int,
                                       caption_text: str,
                                       row2: int = None, cell2: int = None,
                                       title: str = "") -> bool:
    """
    표 셀(table_idx, row, cell)에 이미지 placeholder 텍스트를 삽입.

    - caption 행 (row): ※ 추천 안내문
    - title  행 (row2): < 이미지 제목 >
    """
    try:
        tbl  = doc.tables[table_idx]
        tc   = tbl.rows[row].cells[cell]._tc
        _set_cell_text(tc, caption_text)
        if row2 is not None and cell2 is not None:
            tc2 = tbl.rows[row2].cells[cell2]._tc
            _set_cell_text(tc2, title)
        return True
    except (IndexError, Exception) as _e:
        print(f"    [image_placeholder_cell] 실패: table={table_idx}, "
              f"row={row}, cell={cell}, err={_e}")
        return False


def _insert_image_placeholder_after_para(doc: Document,
                                          keyword: str,
                                          caption_text: str,
                                          title: str = "",
                                          insert_offset: int = 2) -> bool:
    """
    keyword 단락 이후 insert_offset 위치에 이미지 placeholder 블록을 삽입.

    삽입 블록 구조 (3줄):
      ┌─────────────────────────────────────────────────────┐
      │  [이미지 삽입 위치]  <제목>                          │
      │  ※ 추천 안내문 (여러 줄)                             │
      │  ─────────────────── (구분선)                        │
      └─────────────────────────────────────────────────────┘
    """
    body = doc.element.body

    # keyword 단락 탐색
    para = _find_section_para(doc, keyword)
    if para is None:
        return False

    try:
        insert_idx = list(body).index(para._element)
    except ValueError:
        return False

    ref_para = para._element
    insert_pos = insert_idx + insert_offset

    # ── 삽입할 텍스트 줄 목록 ──────────────────────────────────
    header_line = f"━━━ 📸 [이미지 삽입 위치]  {title} ━━━"
    separator   = "─" * 55

    all_lines = (
        [header_line]
        + caption_text.split("\n")
        + [separator]
        # 후행 빈 줄 제거 → collapse_blank_paras 가 정리하므로 불필요
    )

    for offset, line in enumerate(all_lines):
        new_p = copy.deepcopy(ref_para)

        # 모든 텍스트 비우기
        for r in new_p.findall(qn("w:r")):
            for t in r.findall(qn("w:t")): t.text = ""

        # 서식 초기화 (색상·이탤릭 제거)
        for r in new_p.findall(qn("w:r")):
            rpr = r.find(qn("w:rPr"))
            if rpr is not None:
                for tag in [qn("w:color"), qn("w:i"), qn("w:iCs"),
                             qn("w:u"), qn("w:strike")]:
                    el = rpr.find(tag)
                    if el is not None: rpr.remove(el)

        # 들여쓰기·스타일 초기화
        ppr = new_p.find(qn("w:pPr"))
        if ppr is not None:
            for tag in [qn("w:ind"), qn("w:jc"), qn("w:outlineLvl"), qn("w:pStyle")]:
                el = ppr.find(tag)
                if el is not None: ppr.remove(el)
            rpr_ppr = ppr.find(qn("w:rPr"))
            if rpr_ppr is not None:
                for tag in [qn("w:color"), qn("w:i"), qn("w:iCs")]:
                    el = rpr_ppr.find(tag)
                    if el is not None: rpr_ppr.remove(el)

        _write_para(new_p, line)
        body.insert(insert_pos + offset, new_p)

    return True


# ── 이미지 없을 시 빈 테이블(칸) 자동 배정 ───────────────────────────
def _make_blank_image_table(doc: Document,
                             width_emu: int = None,
                             height_emu: int = None,
                             label: str = "[ 이미지 삽입 위치 ]") -> "CT_Tbl":
    """
    지정 크기의 빈 1×1 표를 생성하여 이미지 칸으로 사용.

    Args:
        doc        : 대상 Document
        width_emu  : 표 너비 (EMU 단위). None=페이지 본문 너비 자동 계산
        height_emu : 행 높이 (EMU 단위). None=8cm(2880000 EMU) 기본값
        label      : 빈 칸 안에 표시할 텍스트 (회색 안내문)

    Returns:
        lxml Element (w:tbl)
    """
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement

    if height_emu is None:
        height_emu = int(2880000)      # 8cm

    # 페이지 본문 너비 = 페이지 너비 - 여백
    section = doc.sections[0]
    if width_emu is None:
        page_w   = section.page_width   or int(12240 * 914.4)  # A4
        margin_l = section.left_margin  or int(1800  * 914.4)
        margin_r = section.right_margin or int(1800  * 914.4)
        width_emu = page_w - margin_l - margin_r

    # 표 생성 (python-docx API)
    tbl = doc.add_table(rows=1, cols=1)
    # 'Table Grid' 스타일이 없는 DOCX의 경우 스타일 없이 테두리만 직접 설정
    try:
        tbl.style = 'Table Grid'
    except KeyError:
        pass  # 테두리는 아래 XML 조작으로 처리
    cell = tbl.cell(0, 0)

    # 행 높이 설정
    tr = cell._tc.getparent()
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_emu / 914.4)))   # EMU → twips
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

    # 표 너비 설정
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'),    str(int(width_emu / 914.4)))   # EMU → twips
    tblW.set(qn('w:type'), 'dxa')

    # 셀 테두리: 회색 점선
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        cell._tc.insert(0, tcPr)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'dashed')
        border.set(qn('w:sz'),    '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '808080')
        tcBorders.append(border)
    tcPr.append(tcBorders)

    # 셀 안 세로 중앙 정렬
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

    # 안내 텍스트 삽입 (회색·가운데 정렬)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run  = para.add_run(label)
    run.font.color.rgb = __import__('docx.shared', fromlist=['RGBColor']).RGBColor(0x80, 0x80, 0x80)
    run.font.size  = Pt(11)
    run.font.bold  = False
    run.font.italic = False

    # body에서 표를 분리(detach) 후 반환 → 호출측에서 원하는 위치에 삽입
    tbl_elem = tbl._tbl
    tbl_elem.getparent().remove(tbl_elem)
    return tbl_elem


def insert_blank_image_table_after_para(doc: Document,
                                         keyword: str,
                                         caption_label: str = "[ 이미지 삽입 위치 ]",
                                         height_cm: float = 8.0,
                                         insert_offset: int = 2) -> bool:
    """
    keyword 단락 바로 뒤(insert_offset 위치)에 빈 이미지 칸(표)을 삽입.

    Args:
        doc           : 대상 Document
        keyword       : 기준 단락 검색 텍스트
        caption_label : 빈 칸 안 안내 텍스트
        height_cm     : 칸 높이 (cm, 기본 8)
        insert_offset : keyword 단락으로부터의 삽입 오프셋 (기본 2)

    Returns:
        True: 삽입 성공 / False: keyword 없음 또는 오류
    """
    from docx.shared import Cm

    para = _find_section_para(doc, keyword)
    if para is None:
        return False

    body = doc.element.body
    try:
        insert_idx = list(body).index(para._element)
    except ValueError:
        return False

    height_emu = int(height_cm * 360000)   # cm → EMU (1cm = 360000 EMU)
    tbl_elem   = _make_blank_image_table(doc, height_emu=height_emu,
                                          label=caption_label)
    # 표 삽입 후 구분 단락(빈 줄) 1개 추가
    sep_p = copy.deepcopy(para._element)
    for r in sep_p.findall(qn("w:r")):
        for t in r.findall(qn("w:t")): t.text = ""
    _write_para(sep_p, "")

    body.insert(insert_idx + insert_offset, tbl_elem)
    body.insert(insert_idx + insert_offset + 1, sep_p)
    return True



# ════════════════════════════════════════════════════════════════
# 이미지 실제 삽입 함수 (python-docx 기반)
# ════════════════════════════════════════════════════════════════

def _get_image_drawing_xml(doc: Document, image_path: str, width_cm: float):
    """
    이미지를 doc의 relationship으로 등록하고 w:drawing XML element 반환.
    임시 단락 방식으로 relationship 생성 후 단락만 제거.
    """
    from docx.shared import Cm

    if not os.path.exists(image_path):
        return None, f"파일 없음: {image_path}"

    tmp_p = OxmlElement("w:p")
    doc.element.body.append(tmp_p)

    from docx.text.paragraph import Paragraph as _Para
    tmp_para = _Para(tmp_p, doc)

    try:
        run = tmp_para.add_run()
        run.add_picture(image_path, width=Cm(width_cm))
    except Exception as e:
        doc.element.body.remove(tmp_p)
        return None, str(e)

    drawings = tmp_p.findall(".//" + qn("w:drawing"))
    if not drawings:
        doc.element.body.remove(tmp_p)
        return None, "drawing 요소를 찾을 수 없음"

    drawing_xml = copy.deepcopy(drawings[0])
    doc.element.body.remove(tmp_p)
    return drawing_xml, None


def insert_image_in_cell(doc: Document,
                          table_idx: int,
                          row: int,
                          cell_idx: int,
                          image_path: str,
                          width_cm: float = 12.0,
                          caption: str = None) -> bool:
    """표 셀에 이미지 삽입. 기존 내용 삭제 후 이미지 run 삽입."""
    try:
        tbl  = doc.tables[table_idx]
        cell = tbl.rows[row].cells[cell_idx]
    except (IndexError, Exception):
        return False

    drawing_xml, err = _get_image_drawing_xml(doc, image_path, width_cm)
    if err:
        return False

    tc = cell._tc
    paras = tc.findall(qn("w:p"))
    for p in paras[1:]:
        tc.remove(p)
    for t in paras[0].findall(".//" + qn("w:t")):
        t.text = ""
    for r in paras[0].findall(qn("w:r")):
        paras[0].remove(r)

    img_r = OxmlElement("w:r")
    img_r.append(drawing_xml)
    paras[0].append(img_r)

    if caption:
        cap_p = OxmlElement("w:p")
        cap_r = OxmlElement("w:r")
        cap_rpr = OxmlElement("w:rPr")
        sz  = OxmlElement("w:sz");   sz.set(qn("w:val"), "18"); cap_rpr.append(sz)
        col = OxmlElement("w:color"); col.set(qn("w:val"), "666666"); cap_rpr.append(col)
        cap_r.append(cap_rpr)
        cap_t = OxmlElement("w:t");  cap_t.text = caption
        cap_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        cap_r.append(cap_t);  cap_p.append(cap_r);  tc.append(cap_p)
    return True


def insert_image_after_para(doc: Document,
                             keyword: str,
                             image_path: str,
                             width_cm: float = 14.0,
                             caption: str = None,
                             insert_offset: int = 1) -> bool:
    """keyword 단락 이후(insert_offset 위치)에 이미지 단락 삽입."""
    drawing_xml, err = _get_image_drawing_xml(doc, image_path, width_cm)
    if err:
        return False

    para = _find_section_para(doc, keyword)
    if para is None:
        return False

    body = doc.element.body
    try:
        base_idx = list(body).index(para._element)
    except ValueError:
        return False

    insert_idx = base_idx + insert_offset

    img_p = OxmlElement("w:p")
    ppr   = OxmlElement("w:pPr")
    jc    = OxmlElement("w:jc"); jc.set(qn("w:val"), "center"); ppr.append(jc)
    img_p.append(ppr)
    img_r = OxmlElement("w:r")
    img_r.append(drawing_xml)
    img_p.append(img_r)
    body.insert(insert_idx, img_p)

    if caption:
        cap_p   = OxmlElement("w:p")
        cap_ppr = OxmlElement("w:pPr")
        cap_jc  = OxmlElement("w:jc"); cap_jc.set(qn("w:val"), "center"); cap_ppr.append(cap_jc)
        cap_p.append(cap_ppr)
        cap_r   = OxmlElement("w:r")
        cap_rpr = OxmlElement("w:rPr")
        sz  = OxmlElement("w:sz");    sz.set(qn("w:val"), "18"); cap_rpr.append(sz)
        col = OxmlElement("w:color"); col.set(qn("w:val"), "666666"); cap_rpr.append(col)
        cap_r.append(cap_rpr)
        cap_t = OxmlElement("w:t"); cap_t.text = caption
        cap_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        cap_r.append(cap_t); cap_p.append(cap_r)
        body.insert(insert_idx + 1, cap_p)
    return True


def trim_table_empty_rows(doc: Document, table_idx: int,
                           data_start_row: int = 1) -> int:
    """
    table_idx 표에서 data_start_row 이후의 완전히 빈 행(모든 셀 텍스트가 공백)을
    삭제하여 데이터 행만 남긴다.

    Args:
        doc:            python-docx Document
        table_idx:      삭제 대상 표 인덱스
        data_start_row: 헤더 다음 첫 데이터 행 번호 (기본 1)

    Returns:
        삭제된 행 수
    """
    try:
        tbl = doc.tables[table_idx]
    except IndexError:
        return 0

    tbl_elem = tbl._tbl
    rows = tbl_elem.findall(qn("w:tr"))
    deleted = 0

    # data_start_row 이후 행들 중 모든 셀이 비어있는 행 제거
    for tr in rows[data_start_row:]:
        tcs = tr.findall(qn("w:tc"))
        all_empty = all(
            not "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip()
            for tc in tcs
        )
        if all_empty:
            tbl_elem.remove(tr)
            deleted += 1

    return deleted


# ── 메인 함수 ────────────────────────────────────────────────────
def write_docx(template_path: str, output_path: str,
               injector_content: dict,
               delete_guides: bool = True,
               max_blank_lines: int = 1) -> dict:
    """
    injector_content를 template DOCX에 채워 output_path로 저장.

    실행 순서 (중요):
      1. 셀 채우기           ← TABLE_INDEX 원본 기준 유효
      2. 섹션 삽입           ← 단락 키워드 탐색 후 바로 뒤에 삽입
      2-B. 이미지 Placeholder ← 이미지 없는 위치에 추천 텍스트 삽입
      3. 작성요령 삭제       ← 삽입 완료 후 삭제 (순서 뒤바뀌면 표 사라짐!)
      4. 연속 공백 압축      ← 삽입으로 생긴 여백 정리

    Args:
        template_path:    원본 양식 경로
        output_path:      출력 경로
        injector_content: {
            "cells":    [...],
            "sections": [...],
            "image_placeholders": {   # image_advisor가 생성, 선택적
                slot_id: {
                    "location":     "table_cell" | "after_paragraph",
                    "table_index":  int,
                    "row":  int,  "cell":  int,
                    "row2": int,  "cell2": int,   # title 행 (table_cell 전용)
                    "para_keyword": str,           # after_paragraph 전용
                    "insert_offset": int,
                    "title":        str,
                    "caption_text": str,
                }
            }
        }
        delete_guides:    True이면 작성요령 표·파란색 단락 삭제
        max_blank_lines:  연속 빈 단락 허용 최대수 (기본 1)
    """
    shutil.copy2(template_path, output_path)
    doc = Document(output_path)

    cells_written        = 0
    sections_written     = 0
    images_placeholder   = 0
    errors               = []

    tables = doc.tables

    # ── STEP 1: 셀 채우기 ───────────────────────────────────────
    for item in injector_content.get("cells", []):
        tbl_idx  = item["table"]
        row_idx  = item["row"]
        cell_idx = item["cell"]
        text     = item.get("text", "")
        if not text: continue
        try:
            cell = tables[tbl_idx].rows[row_idx].cells[cell_idx]
            _set_cell_text(cell._tc, text)
            cells_written += 1
        except (IndexError, Exception) as e:
            errors.append(f"cell[{tbl_idx},{row_idx},{cell_idx}]: {e}")

    # ── STEP 2: 섹션 삽입 ───────────────────────────────────────
    # ── STEP 1-B: 표 잉여 빈 행 정리 (데이터 삽입 후 남은 빈 행 제거) ──
    # trim_tables: {table_idx: data_start_row} 형식으로 전달
    empty_rows_removed = 0
    for item in injector_content.get("trim_tables", []):
        tbl_idx        = item.get("table_index")
        data_start_row = item.get("data_start_row", 1)
        if tbl_idx is None:
            continue
        n = trim_table_empty_rows(doc, tbl_idx, data_start_row)
        empty_rows_removed += n


    # ── STEP 1-C: 실제 이미지 삽입 ─────────────────────────────
    # image_slots: {slot_id: {location, image_path, table_index, row, cell,
    #                         para_keyword, insert_offset, width_cm, caption}}
    images_inserted = 0
    image_slots = injector_content.get("image_slots", {})
    for slot_id, slot in image_slots.items():
        img_path = slot.get("image_path", "")
        if not img_path or not os.path.exists(img_path):
            errors.append(f"이미지 슬롯 파일 없음: {img_path!r} ({slot_id})")
            continue
        location  = slot.get("location", "after_paragraph")
        width_cm  = slot.get("width_cm", 12.0)
        caption   = slot.get("caption", "")
        if location == "table_cell":
            ok = insert_image_in_cell(
                doc,
                table_idx = slot.get("table_index", 3),
                row       = slot.get("row", 6),
                cell_idx  = slot.get("cell", 2),
                image_path= img_path,
                width_cm  = width_cm,
                caption   = caption,
            )
        elif location == "after_paragraph":
            ok = insert_image_after_para(
                doc,
                keyword      = slot.get("para_keyword", ""),
                image_path   = img_path,
                width_cm     = width_cm,
                caption      = caption,
                insert_offset= slot.get("insert_offset", 1),
            )
        else:
            ok = False
        if ok:
            images_inserted += 1
        else:
            errors.append(f"이미지 삽입 실패: {slot_id}")

    # ── STEP 2: 섹션 삽입 ───────────────────────────────────────
    for item in injector_content.get("sections", []):
        keyword = item["keyword"]
        text    = item.get("text", "")
        if not text: continue
        para = _find_section_para(doc, keyword)
        if para is None:
            errors.append(f"keyword 단락 없음: '{keyword}'")
            continue
        ok = _insert_text_after_para(para, text, doc)
        if ok:
            sections_written += 1
        else:
            errors.append(f"삽입 위치 없음: '{keyword}'")

    # ── STEP 2-A: 본문 실제 표 삽입 ────────────────────────────
    tables_inserted = 0
    for item in injector_content.get("paragraph_tables", []):
        ok = insert_table_after_para(
            doc,
            keyword=item.get("keyword", ""),
            title=item.get("title", ""),
            headers=item.get("headers", []),
            rows=item.get("rows", []),
            note=item.get("note", ""),
            insert_offset=item.get("insert_offset", 1),
            column_widths_cm=item.get("column_widths_cm", []),
        )
        if ok:
            tables_inserted += 1
        else:
            errors.append(f"본문 표 삽입 실패: keyword={item.get('keyword', '')!r}")

    # ── STEP 2-B: 이미지 Placeholder 삽입 ──────────────────────
    img_placeholders = injector_content.get("image_placeholders", {})
    for slot_id, ph in img_placeholders.items():
        location     = ph.get("location", "table_cell")
        caption_text = ph.get("caption_text", "")
        title        = ph.get("title", "")

        if location == "table_cell":
            ok = _insert_image_placeholder_in_cell(
                doc,
                table_idx = ph.get("table_index", 3),
                row       = ph.get("row",  6),
                cell      = ph.get("cell", 2),
                caption_text = caption_text,
                row2      = ph.get("row2"),
                cell2     = ph.get("cell2"),
                title     = title,
            )
        elif location == "after_paragraph":
            ok = _insert_image_placeholder_after_para(
                doc,
                keyword      = ph.get("para_keyword", ""),
                caption_text = caption_text,
                title        = title,
                insert_offset= ph.get("insert_offset", 2),
            )
        else:
            ok = False

        if ok:
            images_placeholder += 1
        else:
            errors.append(f"image placeholder 삽입 실패: {slot_id}")

    # ── STEP 2-C: 이미지 없을 시 빈 칸(Table) 자동 배정 ────────
    blank_img_slots = injector_content.get("blank_image_slots", [])
    blank_tables_inserted = 0
    for bslot in blank_img_slots:
        keyword      = bslot.get("para_keyword", "")
        label        = bslot.get("label", "[ 이미지 삽입 위치 ]")
        height_cm    = bslot.get("height_cm", 8.0)
        offset       = bslot.get("insert_offset", 2)
        ok = insert_blank_image_table_after_para(
            doc, keyword=keyword, caption_label=label,
            height_cm=height_cm, insert_offset=offset
        )
        if ok:
            blank_tables_inserted += 1
        else:
            errors.append(f"blank image table 삽입 실패: keyword={keyword!r}")

    # ── STEP 3: 작성요령 삭제 (삽입 완료 후) ───────────────────
    guide_stats = {"tables_deleted": 0, "paras_deleted": 0, "blue_rows_deleted": 0}
    if delete_guides:
        guide_stats = delete_guide_elements(doc)

    # ── STEP 4: 연속 공백 단락 압축 ────────────────────────────
    blank_removed = collapse_blank_paras(doc, max_blank=max_blank_lines)

    doc.save(output_path)

    return {
        "cells_written":          cells_written,
        "sections_written":       sections_written,
        "images_inserted":         images_inserted,
        "images_placeholder":     images_placeholder,
        "tables_inserted":        tables_inserted,
        "blank_tables_inserted":  blank_tables_inserted,
        "tables_deleted":         guide_stats["tables_deleted"],
        "paras_deleted":          guide_stats["paras_deleted"],
        "blue_rows_deleted":      guide_stats.get("blue_rows_deleted", 0),
        "blank_removed":          blank_removed,
        "empty_rows_removed":     empty_rows_removed,
        "errors":                 errors,
        "output":                 output_path,
    }
